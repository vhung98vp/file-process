import os
import csv
import pandas as pd
from docx import Document
import fitz
from .utils import is_copyable_pdf_page, process_copyable_pdf_page


def read_docx(path):
    doc = Document(path)
    text_lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_lines.append(p.text.strip())
    texts = [{"page": 0,  "text": text_lines}] if text_lines else []

    tables = []
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_data.append(row_text)
        tables.append({"page": 0, "table": table_idx, "cells": table_data})
    
    return texts, tables


def read_txt(path):
    with open(path, encoding="utf-8", errors="ignore") as f:
        text_lines = [line.strip() for line in f if line.strip()]
    texts = [{"page": 0,  "text": text_lines}] if text_lines else []

    return texts, []


def read_xlsx(path):
    tables = []
    try:
        xls = pd.ExcelFile(path)
        for page_idx, sheet_name in enumerate(xls.sheet_names):
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
            df = df.dropna(how="all")
            if not df.empty:
                cells = df.fillna("").astype(str).values.tolist()
                tables.append({"page": page_idx, "table": 0, "cells": cells})
    except Exception as e:
        print(f"Error reading {path}: {e}")
    return [], tables


def read_csv(path):
    text_lines = []
    cells = []
    with open(path, encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            if len(row) == 1:
                text_lines.append(row[0])
            else:
                cells.append(row)
    
    texts = [{"page": 0,  "text": text_lines}] if text_lines else []
    tables = [{"page": 0, "table": 0, "cells": cells}] if cells else []
    return texts, tables


def read_pdf(path):
    texts, tables = [], []
    with fitz.open(path) as pages:
        first_page = pages[0] if len(pages) > 0 else None
        if first_page and is_copyable_pdf_page(first_page):
            for page_idx, page in enumerate(pages):
                p_texts, p_tables = process_copyable_pdf_page(page, page_idx)
                texts.extend(p_texts)
                tables.extend(p_tables)

    return texts, tables

